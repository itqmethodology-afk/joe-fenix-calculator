"""
Funzioni di utility per il Joe Fenix Method Calculator
"""

import pandas as pd
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from config import *


def calcola_bmr(peso, altezza, eta, sesso):
    """
    Calcola il Metabolismo Basale (BMR) usando la formula Mifflin-St Jeor
    
    Args:
        peso (float): Peso in kg
        altezza (int): Altezza in cm
        eta (int): Età in anni
        sesso (str): "Maschio" o "Femmina"
    
    Returns:
        float: BMR in kcal/giorno
    """
    if sesso == "Maschio":
        return 10 * peso + 6.25 * altezza - 5 * eta + 5
    else:
        return 10 * peso + 6.25 * altezza - 5 * eta - 161


def calcola_tdee(bmr, livello_attivita):
    """
    Calcola il Total Daily Energy Expenditure (TDEE)
    
    Args:
        bmr (float): Metabolismo basale
        livello_attivita (str): Livello di attività fisica
    
    Returns:
        float: TDEE in kcal/giorno
    """
    return bmr * FATTORI_ATTIVITA[livello_attivita]


def calcola_calorie_target(tdee, obiettivo):
    """
    Calcola le calorie target in base all'obiettivo
    
    Args:
        tdee (float): Total Daily Energy Expenditure
        obiettivo (str): "Perdere peso", "Mantenere peso", "Aumentare massa"
    
    Returns:
        float: Calorie target giornaliere
    """
    if obiettivo == "Perdere peso":
        return tdee - DEFICIT_PERDITA_PESO
    elif obiettivo == "Aumentare massa":
        return tdee + SURPLUS_MASSA
    else:
        return tdee


def calcola_macronutrienti(peso, calorie_target):
    """
    Calcola la distribuzione dei macronutrienti
    
    Args:
        peso (float): Peso corporeo in kg
        calorie_target (float): Calorie giornaliere target
    
    Returns:
        dict: Dizionario con proteine, grassi e carboidrati in grammi
    """
    # Proteine: 2.2g per kg (high protein approach)
    proteine_g = peso * PROTEINE_PER_KG
    
    # Grassi: 25% delle calorie totali
    grassi_g = (calorie_target * PERCENTUALE_GRASSI) / CALORIE_PER_G_GRASSI
    
    # Carboidrati: calorie rimanenti
    carboidrati_g = (calorie_target - (proteine_g * CALORIE_PER_G_PROTEINE) - (grassi_g * CALORIE_PER_G_GRASSI)) / CALORIE_PER_G_CARBOIDRATI
    
    return {
        "proteine": round(proteine_g, 1),
        "grassi": round(grassi_g, 1),
        "carboidrati": round(carboidrati_g, 1)
    }


def calcola_bmi(peso, altezza):
    """
    Calcola il Body Mass Index (BMI)
    
    Args:
        peso (float): Peso in kg
        altezza (int): Altezza in cm
    
    Returns:
        float: BMI
    """
    altezza_m = altezza / 100
    return peso / (altezza_m ** 2)


def interpreta_bmi(bmi):
    """
    Interpreta il valore del BMI
    
    Args:
        bmi (float): Body Mass Index
    
    Returns:
        str: Categoria BMI
    """
    if bmi < 18.5:
        return "Sottopeso"
    elif bmi < 25:
        return "Normopeso"
    elif bmi < 30:
        return "Sovrappeso"
    else:
        return "Obesità"


def genera_piano_settimanale(macros, dieta="Onnivoro"):
    """
    Genera un piano settimanale base
    
    Args:
        macros (dict): Macronutrienti giornalieri
        dieta (str): Tipo di dieta
    
    Returns:
        pd.DataFrame: Piano settimanale
    """
    giorni = ["Lunedì", "Martedì", "Mercoledì", "Giovedì", "Venerdì", "Sabato", "Domenica"]
    
    piano = []
    for giorno in giorni:
        piano.append({
            "Giorno": giorno,
            "Proteine (g)": macros["proteine"],
            "Carboidrati (g)": macros["carboidrati"],
            "Grassi (g)": macros["grassi"],
            "Calorie": round(
                macros["proteine"] * 4 + 
                macros["carboidrati"] * 4 + 
                macros["grassi"] * 9
            )
        })
    
    return pd.DataFrame(piano)


def genera_lista_spesa(dieta="Onnivoro", intolleranze=[]):
    """
    Genera una lista della spesa intelligente
    
    Args:
        dieta (str): Tipo di dieta
        intolleranze (list): Lista di intolleranze
    
    Returns:
        dict: Lista della spesa per categoria
    """
    lista = {
        "Proteine": GROCERY_PROTEINE.copy(),
        "Carboidrati": GROCERY_CARBOIDRATI.copy(),
        "Grassi": GROCERY_GRASSI.copy(),
        "Verdure": GROCERY_VERDURE.copy()
    }
    
    # Filtra in base a dieta
    if dieta == "Vegano":
        lista["Proteine"] = ["Legumi (lenticchie, ceci)", "Tofu", "Tempeh", "Seitan"]
    elif dieta == "Vegetariano":
        lista["Proteine"] = ["Uova", "Yogurt greco", "Legumi (lenticchie, ceci)", "Tofu"]
    
    # Filtra in base a intolleranze
    if "Lattosio" in intolleranze:
        lista["Proteine"] = [p for p in lista["Proteine"] if "yogurt" not in p.lower()]
    
    if "Glutine" in intolleranze:
        lista["Carboidrati"] = [c for c in lista["Carboidrati"] if "avena" not in c.lower()]
    
    if "Uova" in intolleranze:
        lista["Proteine"] = [p for p in lista["Proteine"] if "uova" not in p.lower()]
    
    if "Frutta secca" in intolleranze:
        lista["Grassi"] = [g for g in lista["Grassi"] if "frutta secca" not in g.lower() and "mandorle" not in g.lower() and "noci" not in g.lower()]
    
    return lista


def esporta_pdf(dati_utente, macros, piano_settimanale, filename="joe_fenix_plan.pdf"):
    """
    Esporta il piano in formato PDF
    
    Args:
        dati_utente (dict): Dati personali dell'utente
        macros (dict): Macronutrienti
        piano_settimanale (pd.DataFrame): Piano settimanale
        filename (str): Nome del file
    
    Returns:
        str: Path del file generato
    """
    doc = SimpleDocTemplate(filename, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()
    
    # Titolo
    title = Paragraph("🔥 Joe Fenix Method - Il Tuo Piano Personalizzato", styles['Title'])
    elements.append(title)
    elements.append(Spacer(1, 12))
    
    # Dati utente
    data_utente = [
        ["Età", f"{dati_utente['eta']} anni"],
        ["Peso", f"{dati_utente['peso']} kg"],
        ["Altezza", f"{dati_utente['altezza']} cm"],
        ["BMI", f"{dati_utente['bmi']:.1f}"]
    ]
    
    table_utente = Table(data_utente)
    table_utente.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    elements.append(table_utente)
    elements.append(Spacer(1, 12))
    
    # Macronutrienti
    subtitle = Paragraph("Target Giornalieri", styles['Heading2'])
    elements.append(subtitle)
    
    data_macros = [
        ["Calorie", "Proteine", "Carboidrati", "Grassi"],
        [
            f"{dati_utente['calorie']} kcal",
            f"{macros['proteine']} g",
            f"{macros['carboidrati']} g",
            f"{macros['grassi']} g"
        ]
    ]
    
    table_macros = Table(data_macros)
    table_macros.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    elements.append(table_macros)
    
    doc.build(elements)
    return filename


def esporta_word(dati_utente, macros, lista_spesa, filename="joe_fenix_plan.docx"):
    """
    Esporta il piano in formato Word
    
    Args:
        dati_utente (dict): Dati personali
        macros (dict): Macronutrienti
        lista_spesa (dict): Lista della spesa
        filename (str): Nome del file
    
    Returns:
        str: Path del file generato
    """
    doc = Document()
    
    # Titolo
    title = doc.add_heading('🔥 Joe Fenix Method - Il Tuo Piano', 0)
    title.alignment = 1  # Centro
    
    # Dati utente
    doc.add_heading('I Tuoi Dati', level=1)
    doc.add_paragraph(f"Età: {dati_utente['eta']} anni")
    doc.add_paragraph(f"Peso: {dati_utente['peso']} kg")
    doc.add_paragraph(f"Altezza: {dati_utente['altezza']} cm")
    doc.add_paragraph(f"BMI: {dati_utente['bmi']:.1f}")
    
    # Macronutrienti
    doc.add_heading('Target Giornalieri', level=1)
    doc.add_paragraph(f"Calorie: {dati_utente['calorie']} kcal")
    doc.add_paragraph(f"Proteine: {macros['proteine']} g")
    doc.add_paragraph(f"Carboidrati: {macros['carboidrati']} g")
    doc.add_paragraph(f"Grassi: {macros['grassi']} g")
    
    # Lista spesa
    doc.add_heading('Smart Grocery List', level=1)
    for categoria, items in lista_spesa.items():
        doc.add_heading(categoria, level=2)
        for item in items:
            doc.add_paragraph(f"• {item}", style='List Bullet')
    
    doc.save(filename)
    return filename
