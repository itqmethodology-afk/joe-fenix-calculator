"""
Export utilities for PDF, Word, and Web formats
"""

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.units import cm
from docx import Document
from docx.shared import RGBColor, Pt
import io

# Brand colors
SAGE_GREEN = (139, 157, 131)
ENERGY_ORANGE = (255, 107, 53)

def export_to_pdf(user_data, macros, weekly_plan, grocery_list):
    """Generate PDF report"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    story = []
    styles = getSampleStyleSheet()
    
    # Custom title style
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.Color(SAGE_GREEN[0]/255, SAGE_GREEN[1]/255, SAGE_GREEN[2]/255),
        spaceAfter=30
    )
    
    # Title
    story.append(Paragraph("Joe Fenix Method - Your Personalized Plan", title_style))
    story.append(Spacer(1, 0.5*cm))
    
    # User info
    story.append(Paragraph(f"<b>Name:</b> {user_data.get('name', 'N/A')}", styles['Normal']))
    story.append(Paragraph(f"<b>Age:</b> {user_data['age']} | <b>Weight:</b> {user_data['weight']}kg | <b>Height:</b> {user_data['height']}cm", styles['Normal']))
    story.append(Paragraph(f"<b>BMI:</b> {user_data['bmi']}", styles['Normal']))
    story.append(Spacer(1, 0.5*cm))
    
    # Daily targets
    story.append(Paragraph("<b>Your Daily Targets</b>", styles['Heading2']))
    targets_data = [
        ['Calories', 'Protein', 'Carbs', 'Fat'],
        [f"{macros['calories']}", f"{macros['protein_g']}g", f"{macros['carbs_g']}g", f"{macros['fat_g']}g"]
    ]
    targets_table = Table(targets_data)
    targets_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.Color(SAGE_GREEN[0]/255, SAGE_GREEN[1]/255, SAGE_GREEN[2]/255)),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey)
    ]))
    story.append(targets_table)
    story.append(Spacer(1, 0.5*cm))
    
    # The Perfect Day
    story.append(Paragraph("<b>The Perfect Day - Protein Distribution</b>", styles['Heading2']))
    for meal, protein in weekly_plan['meal_distribution'].items():
        story.append(Paragraph(f"• <b>{meal.capitalize()}:</b> {protein}", styles['Normal']))
    story.append(Spacer(1, 0.5*cm))
    
    # Grocery list
    story.append(Paragraph("<b>Smart Grocery List</b>", styles['Heading2']))
    for category, items in grocery_list.items():
        story.append(Paragraph(f"<b>{category}:</b>", styles['Normal']))
        for item in items:
            story.append(Paragraph(f"  • {item}", styles['Normal']))
        story.append(Spacer(1, 0.3*cm))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def export_to_word(user_data, macros, weekly_plan, grocery_list):
    """Generate Word document"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Joe Fenix Method - Your Personalized Plan', 0)
    title.runs[0].font.color.rgb = RGBColor(*SAGE_GREEN)
    
    # User info
    doc.add_paragraph(f"Name: {user_data.get('name', 'N/A')}")
    doc.add_paragraph(f"Age: {user_data['age']} | Weight: {user_data['weight']}kg | Height: {user_data['height']}cm")
    doc.add_paragraph(f"BMI: {user_data['bmi']}")
    doc.add_paragraph()
    
    # Daily targets
    doc.add_heading('Your Daily Targets', 2)
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Light Grid Accent 1'
    headers = table.rows[0].cells
    headers[0].text = 'Calories'
    headers[1].text = 'Protein'
    headers[2].text = 'Carbs'
    headers[3].text = 'Fat'
    
    values = table.rows[1].cells
    values[0].text = str(macros['calories'])
    values[1].text = f"{macros['protein_g']}g"
    values[2].text = f"{macros['carbs_g']}g"
    values[3].text = f"{macros['fat_g']}g"
    doc.add_paragraph()
    
    # The Perfect Day
    doc.add_heading('The Perfect Day - Protein Distribution', 2)
    for meal, protein in weekly_plan['meal_distribution'].items():
        doc.add_paragraph(f"{meal.capitalize()}: {protein}", style='List Bullet')
    doc.add_paragraph()
    
    # Grocery list
    doc.add_heading('Smart Grocery List', 2)
    for category, items in grocery_list.items():
        doc.add_heading(category, 3)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_web_view(user_data, macros, weekly_plan, grocery_list):
    """Generate HTML for web view"""
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Joe Fenix Method - Your Plan</title>
        <style>
            body {{
                font-family: Arial, sans-serif;
                max-width: 800px;
                margin: 0 auto;
                padding: 20px;
                background-color: 
;
            }}
            h1 {{
                color: #{SAGE_GREEN[0]:02x}{SAGE_GREEN[1]:02x}{SAGE_GREEN[2]:02x};
            }}
            .targets {{
                background-color: white;
                padding: 20px;
                border-radius: 10px;
                margin: 20px 0;
            }}
            .metric {{
                display: inline-block;
                margin: 10px 20px;
                text-align: center;
            }}
            .metric-value {{
                font-size: 24px;
                font-weight: bold;
                color: #{ENERGY_ORANGE[0]:02x}{ENERGY_ORANGE[1]:02x}{ENERGY_ORANGE[2]:02x};
            }}
            ul {{
                line-height: 1.8;
            }}
        </style>
    </head>
    <body>
        <h1>🔥 Joe Fenix Method - Your Personalized Plan</h1>
        <p><strong>BUILD. PROTECT. FUEL.</strong></p>
        
        <div class="targets">
            <h2>Your Daily Targets</h2>
            <div class="metric">
                <div class="metric-value">{macros['calories']}</div>
                <div>Calories</div>
            </div>
            <div class="metric">
                <div class="metric-value">{macros['protein_g']}g</div>
                <div>Protein</div>
            </div>
            <div class="metric">
                <div class="metric-value">{macros['carbs_g']}g</div>
                <div>Carbs</div>
            </div>
            <div class="metric">
                <div class="metric-value">{macros['fat_g']}g</div>
                <div>Fat</div>
            </div>
        </div>
        
        <h2>The Perfect Day</h2>
        <ul>
    """
    
    for meal, protein in weekly_plan['meal_distribution'].items():
        html += f"<li><strong>{meal.capitalize()}:</strong> {protein}</li>"
    
    html += """
        </ul>
        
        <h2>Smart Grocery List</h2>
    """
    
    for category, items in grocery_list.items():
        html += f"<h3>{category}</h3><ul>"
        for item in items:
            html += f"<li>{item}</li>"
        html += "</ul>"
    
    html += """
    </body>
    </html>
    """
    
    return html
